"""
build_docx.py — Converts Class 11 Chapter 4 (Animal Kingdom) into an editable Microsoft Word (.docx) document.
Preserves 100% of the content, structure, tables, callout blocks, and figures.
"""

import os
import sys
import html
import re

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

HERE = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(HERE, "assets")
OUT_DOCX = os.path.join(HERE, "Ch4_AnimalKingdom.docx")
SRC_PY = os.path.join(HERE, "Ch4_AnimalKingdom.py")


def add_formatted_runs(paragraph, html_str, base_bold=False, base_italic=False,
                       font_name="Calibri", font_size=Pt(10.5), font_color=RGBColor(0x1F, 0x29, 0x37)):
    """Parses simple inline HTML tags (<b>, <i>, <br>) and appends formatted runs."""
    tokens = re.split(r'(</?[bi]>|<br\s*/?>)', html_str)
    bold_state = base_bold
    italic_state = base_italic

    for token in tokens:
        if not token:
            continue
        t_low = token.lower()
        if t_low == "<b>":
            bold_state = True
        elif t_low == "</b>":
            bold_state = False
        elif t_low == "<i>":
            italic_state = True
        elif t_low == "</i>":
            italic_state = False
        elif t_low.startswith("<br"):
            paragraph.add_run("\n")
        else:
            clean = html.unescape(token)
            if clean:
                run = paragraph.add_run(clean)
                run.bold = bold_state
                run.italic = italic_state
                if font_name:
                    run.font.name = font_name
                if font_size:
                    run.font.size = font_size
                if font_color:
                    run.font.color.rgb = font_color


def generate_docx():
    doc = docx.Document()

    # 1. Page Margins & Setup (A4 size with clean 0.7 in margins)
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)

    # 2. Base Normal Style
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(10.5)
    normal_font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(3.5)
    normal_style.paragraph_format.space_before = Pt(0)

    # 3. Builder Helper Functions
    def title_block(title_text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.87)

        # Slate header fill
        shading = parse_xml(r'<w:shd {} w:fill="0F172A"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)

        tcMar = parse_xml(r'<w:tcMar {}><w:top w:w="220" w:type="dxa"/><w:bottom w:w="220" w:type="dxa"/><w:left w:w="280" w:type="dxa"/><w:right w:w="280" w:type="dxa"/></w:tcMar>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(tcMar)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

        r_sub = p.add_run("CLASS 11 BIOLOGY • CHAPTER 4\n")
        r_sub.font.name = "Calibri"
        r_sub.font.size = Pt(10)
        r_sub.font.bold = True
        r_sub.font.color.rgb = RGBColor(0x38, 0xBD, 0xF8)

        r_title = p.add_run(title_text.upper() + "\n")
        r_title.font.name = "Calibri"
        r_title.font.size = Pt(20)
        r_title.font.bold = True
        r_title.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        r_desc = p.add_run("Complete NEET Revision Notes • NCERT-Pinned Figures & High-Yield Summary")
        r_desc.font.name = "Calibri"
        r_desc.font.size = Pt(9.5)
        r_desc.font.italic = True
        r_desc.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after = Pt(4)
        return []

    def heading(number, text, level=1, has_table=False):
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True

        if level == 1:
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            r_num = p.add_run(f"{number}  ")
            r_num.font.name = "Calibri"
            r_num.font.size = Pt(14)
            r_num.font.bold = True
            r_num.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)

            r_text = p.add_run(text.upper())
            r_text.font.name = "Calibri"
            r_text.font.size = Pt(14)
            r_text.font.bold = True
            r_text.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

            pBdr = parse_xml(r'<w:pBdr {}><w:bottom w:val="single" w:sz="8" w:space="2" w:color="0284C7"/></w:pBdr>'.format(nsdecls('w')))
            p._p.get_or_add_pPr().append(pBdr)

        elif level == 2:
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            r_num = p.add_run(f"{number}  ")
            r_num.font.name = "Calibri"
            r_num.font.size = Pt(12)
            r_num.font.bold = True
            r_num.font.color.rgb = RGBColor(0x03, 0x69, 0xA1)

            r_text = p.add_run(text)
            r_text.font.name = "Calibri"
            r_text.font.size = Pt(12)
            r_text.font.bold = True
            r_text.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

        else:
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            r_num = p.add_run(f"{number}  ")
            r_num.font.name = "Calibri"
            r_num.font.size = Pt(11)
            r_num.font.bold = True
            r_num.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)

            r_text = p.add_run(text)
            r_text.font.name = "Calibri"
            r_text.font.size = Pt(11)
            r_text.font.bold = True
            r_text.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

        return p

    def body(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3.5)
        add_formatted_runs(p, text)
        return p

    def b1(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2.5)
        p.paragraph_format.left_indent = Inches(0.25)
        clean = text.removeprefix('&bull; ').strip()
        add_formatted_runs(p, clean)
        return p

    def b2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.45)
        r_bullet = p.add_run("–  ")
        r_bullet.bold = True
        r_bullet.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        clean = text.removeprefix('- ').strip()
        add_formatted_runs(p, clean)
        return p

    def keyterm(text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.87)

        shading = parse_xml(r'<w:shd {} w:fill="F0F9FF"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)
        borders = parse_xml(r'''
            <w:tcBorders {} >
                <w:left w:val="single" w:sz="18" w:space="0" w:color="0284C7"/>
                <w:top w:val="single" w:sz="4" w:space="0" w:color="E0F2FE"/>
                <w:right w:val="single" w:sz="4" w:space="0" w:color="E0F2FE"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="E0F2FE"/>
            </w:tcBorders>
        '''.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(borders)
        tcMar = parse_xml(r'<w:tcMar {}><w:top w:w="80" w:type="dxa"/><w:bottom w:w="80" w:type="dxa"/><w:left w:w="140" w:type="dxa"/><w:right w:w="140" w:type="dxa"/></w:tcMar>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(tcMar)

        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r_lbl = p.add_run("KEY TERM • ")
        r_lbl.font.name = "Calibri"
        r_lbl.font.size = Pt(8.5)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)
        add_formatted_runs(p, text)

        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after = Pt(2)
        return tbl

    def process_flow(steps):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.87)

        shading = parse_xml(r'<w:shd {} w:fill="F8FAFC"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)
        borders = parse_xml(r'''
            <w:tcBorders {} >
                <w:left w:val="single" w:sz="18" w:space="0" w:color="0F172A"/>
                <w:top w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>
                <w:right w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>
            </w:tcBorders>
        '''.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(borders)
        tcMar = parse_xml(r'<w:tcMar {}><w:top w:w="100" w:type="dxa"/><w:bottom w:w="100" w:type="dxa"/><w:left w:w="140" w:type="dxa"/><w:right w:w="140" w:type="dxa"/></w:tcMar>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(tcMar)

        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        r_hd = p.add_run("PROCESS FLOW\n")
        r_hd.font.name = "Calibri"
        r_hd.font.size = Pt(9)
        r_hd.font.bold = True
        r_hd.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

        for i, s in enumerate(steps, 1):
            p_step = cell.add_paragraph()
            p_step.paragraph_format.space_before = Pt(1)
            p_step.paragraph_format.space_after = Pt(1.5)
            r_num = p_step.add_run(f"[{i}]  ")
            r_num.font.bold = True
            r_num.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)
            add_formatted_runs(p_step, s)

        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after = Pt(2)
        return tbl

    def note(text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.87)

        shading = parse_xml(r'<w:shd {} w:fill="FFFBEB"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)
        borders = parse_xml(r'''
            <w:tcBorders {} >
                <w:left w:val="single" w:sz="18" w:space="0" w:color="D97706"/>
                <w:top w:val="single" w:sz="4" w:space="0" w:color="FEF3C7"/>
                <w:right w:val="single" w:sz="4" w:space="0" w:color="FEF3C7"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="FEF3C7"/>
            </w:tcBorders>
        '''.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(borders)
        tcMar = parse_xml(r'<w:tcMar {}><w:top w:w="80" w:type="dxa"/><w:bottom w:w="80" w:type="dxa"/><w:left w:w="140" w:type="dxa"/><w:right w:w="140" w:type="dxa"/></w:tcMar>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(tcMar)

        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r_lbl = p.add_run("NOTE • ")
        r_lbl.font.name = "Calibri"
        r_lbl.font.size = Pt(8.5)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
        add_formatted_runs(p, text)

        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after = Pt(2)
        return tbl

    def memory_aid(text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.87)

        shading = parse_xml(r'<w:shd {} w:fill="F0FDF4"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)
        borders = parse_xml(r'''
            <w:tcBorders {} >
                <w:left w:val="single" w:sz="18" w:space="0" w:color="16A34A"/>
                <w:top w:val="single" w:sz="4" w:space="0" w:color="DCFCE7"/>
                <w:right w:val="single" w:sz="4" w:space="0" w:color="DCFCE7"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="DCFCE7"/>
            </w:tcBorders>
        '''.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(borders)
        tcMar = parse_xml(r'<w:tcMar {}><w:top w:w="80" w:type="dxa"/><w:bottom w:w="80" w:type="dxa"/><w:left w:w="140" w:type="dxa"/><w:right w:w="140" w:type="dxa"/></w:tcMar>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(tcMar)

        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r_lbl = p.add_run("MEMORY AID • ")
        r_lbl.font.name = "Calibri"
        r_lbl.font.size = Pt(8.5)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)
        add_formatted_runs(p, text)

        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after = Pt(2)
        return tbl

    def figure(asset_name, caption_text, max_width_cm=15.9):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(6)
        p_img.paragraph_format.space_after = Pt(2)
        p_img.paragraph_format.keep_with_next = True

        path = os.path.join(ASSETS, asset_name)
        if os.path.exists(path):
            width_in = min(max_width_cm / 2.54, 6.2)
            p_img.add_run().add_picture(path, width=Inches(width_in))

        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(0)
        p_cap.paragraph_format.space_after = Pt(8)
        add_formatted_runs(p_cap, caption_text, font_size=Pt(9),
                           font_color=RGBColor(0x4B, 0x55, 0x63), base_italic=True)
        return p_img

    def figure_pair(left, right, gutter=10):
        """Two figure panels stacked horizontally side-by-side in a 2-column borderless table."""
        tbl = doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        c0 = tbl.cell(0, 0)
        c0.width = Inches(3.35)
        p0_img = c0.paragraphs[0]
        p0_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0_img.paragraph_format.space_before = Pt(4)
        p0_img.paragraph_format.space_after = Pt(2)
        p0_img.paragraph_format.keep_with_next = True

        left_path = os.path.join(ASSETS, left[0])
        if os.path.exists(left_path):
            p0_img.add_run().add_picture(left_path, width=Inches(3.0))

        p0_cap = c0.add_paragraph()
        p0_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0_cap.paragraph_format.space_before = Pt(0)
        p0_cap.paragraph_format.space_after = Pt(4)
        add_formatted_runs(p0_cap, left[1], font_size=Pt(8.5),
                           font_color=RGBColor(0x4B, 0x55, 0x63), base_italic=True)

        c1 = tbl.cell(0, 1)
        c1.width = Inches(3.35)
        p1_img = c1.paragraphs[0]
        p1_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1_img.paragraph_format.space_before = Pt(4)
        p1_img.paragraph_format.space_after = Pt(2)
        p1_img.paragraph_format.keep_with_next = True

        right_path = os.path.join(ASSETS, right[0])
        if os.path.exists(right_path):
            p1_img.add_run().add_picture(right_path, width=Inches(3.0))

        p1_cap = c1.add_paragraph()
        p1_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1_cap.paragraph_format.space_before = Pt(0)
        p1_cap.paragraph_format.space_after = Pt(4)
        add_formatted_runs(p1_cap, right[1], font_size=Pt(8.5),
                           font_color=RGBColor(0x4B, 0x55, 0x63), base_italic=True)

        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after = Pt(6)
        return tbl

    def data_table(rows, col_widths=None, compact=False):
        tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        TOTAL_W = 6.87
        if col_widths:
            total_prop = sum(col_widths)
            widths_in = [w / total_prop * TOTAL_W for w in col_widths]
        else:
            widths_in = [TOTAL_W / len(rows[0])] * len(rows[0])

        fs = Pt(7.5 if compact else 9.0)

        for r_idx, row in enumerate(rows):
            tr = tbl.rows[r_idx]._tr
            trPr = tr.get_or_add_trPr()
            trPr.append(parse_xml(r'<w:cantSplit {}/>'.format(nsdecls('w'))))
            if r_idx == 0:
                trPr.append(parse_xml(r'<w:tblHeader {}/>'.format(nsdecls('w'))))

            is_header = (r_idx == 0)
            bg_color = "1E293B" if is_header else ("F8FAFC" if r_idx % 2 == 1 else "FFFFFF")

            for c_idx, val in enumerate(row):
                cell = tbl.cell(r_idx, c_idx)
                cell.width = Inches(widths_in[c_idx])

                shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), bg_color))
                cell._tc.get_or_add_tcPr().append(shd)

                tcMar = parse_xml(r'<w:tcMar {}><w:top w:w="80" w:type="dxa"/><w:bottom w:w="80" w:type="dxa"/><w:left w:w="90" w:type="dxa"/><w:right w:w="90" w:type="dxa"/></w:tcMar>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(tcMar)

                borders = parse_xml(r'''
                    <w:tcBorders {} >
                        <w:left w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>
                        <w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>
                        <w:right w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>
                        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>
                    </w:tcBorders>
                '''.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(borders)

                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = 1.05
                if is_header:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    add_formatted_runs(p, val, base_bold=True, font_size=Pt(8.0 if compact else 9.5),
                                       font_color=RGBColor(0xFF, 0xFF, 0xFF))
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    add_formatted_runs(p, val, font_size=fs, font_color=RGBColor(0x1F, 0x29, 0x37))

        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after = Pt(6)
        return tbl

    def compact_table(rows, col_widths=None):
        return data_table(rows, col_widths=col_widths, compact=True)

    def PageBreak():
        p = doc.add_paragraph()
        p.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
        return p

    def KeepTogether(items):
        return items

    class DummyStory:
        def __init__(self):
            self.items = []

        def append(self, item):
            self.items.append(item)

        def __iadd__(self, other):
            if isinstance(other, list):
                self.items.extend(other)
            else:
                self.items.append(other)
            return self

    # Execution namespace
    local_env = {
        "title_block": title_block,
        "heading": heading,
        "body": body,
        "b1": b1,
        "b2": b2,
        "keyterm": keyterm,
        "process_flow": process_flow,
        "note": note,
        "memory_aid": memory_aid,
        "figure": figure,
        "figure_pair": figure_pair,
        "data_table": data_table,
        "compact_table": compact_table,
        "PageBreak": PageBreak,
        "KeepTogether": KeepTogether,
        "story": DummyStory(),
    }

    # Read Ch4_AnimalKingdom.py and execute story generation lines
    with open(SRC_PY, "r") as f:
        src = f.read()

    start_token = "story = []"
    end_token = "def main():"
    start_pos = src.find(start_token)
    end_pos = src.find(end_token)
    if start_pos == -1 or end_pos == -1:
        raise RuntimeError("Could not locate story block in Ch4_AnimalKingdom.py")

    exec_code = src[start_pos + len(start_token):end_pos]

    # Execute inside our controlled docx environment
    exec(exec_code, local_env)

    doc.save(OUT_DOCX)
    print(f"Successfully generated docx: {OUT_DOCX}")
    return OUT_DOCX


if __name__ == "__main__":
    generate_docx()
