"""Build an EDITABLE, near-identical Word (.docx) copy of Ch4_AnimalKingdom.pdf.

Goal (per user): a Google-Docs-friendly .docx that, viewed in print layout,
looks as close to the frozen ReportLab PDF as an editable document can — while
staying real, editable Word text/tables (not page images).

How fidelity is preserved without re-typing a single line of content:
  * The chapter's own script (Ch4_AnimalKingdom.py) is executed against a set of
    lightweight shims, so every content string is reused VERBATIM. The shims turn
    each helper call (heading/keyterm/process_flow/note/memory_aid/data_table/
    figure/figure_row/body/b1/b2/title_block) into a small semantic dict.
  * All geometry, fonts, sizes, leading, spacing and colours are read straight
    from the frozen neet_template.py — the same numbers the PDF was built from.
  * Banners, tables, bullets, captions are emitted as native, editable Word
    objects. The few genuinely vector marks (DNA motif, process-flow triangles,
    the NOTE "!" and MEMORY-AID star icons) are rasterised from the SAME vector
    code in neet_template and embedded as small images, so they match exactly.

Run:  python3 Ch4_AnimalKingdom_docx.py
Out:  Ch4_AnimalKingdom.docx  (import into Google Docs losslessly)
"""

import io
import os
import re
import sys

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = HERE
while not os.path.exists(os.path.join(ROOT, "neet_template.py")):
    parent = os.path.dirname(ROOT)
    if parent == ROOT:
        raise RuntimeError("neet_template.py not found in any parent directory")
    ROOT = parent
sys.path.insert(0, ROOT)

ASSETS = os.path.join(HERE, "assets")
CH_SCRIPT = os.path.join(HERE, "Ch4_AnimalKingdom.py")
OUT_DOCX = os.path.join(HERE, "Ch4_AnimalKingdom.docx")
MARKS_DIR = "/tmp/docx_marks_ch4"
os.makedirs(MARKS_DIR, exist_ok=True)

# ======================================================================================
# PHASE 1 — rasterise the vector marks from the REAL neet_template (real reportlab).
# Done first, before any reportlab shim is installed.
# ======================================================================================
import pymupdf  # noqa: E402
from reportlab.graphics import renderPDF  # noqa: E402
from reportlab.pdfbase.pdfmetrics import stringWidth  # noqa: E402
import neet_template as NT  # the real, frozen module  # noqa: E402

CM = 28.3464567  # points per cm
FRAME_WIDTH_PT = float(NT.FRAME_WIDTH)


def _hex(color):
    """ReportLab HexColor -> 'RRGGBB'."""
    return "{:02X}{:02X}{:02X}".format(
        int(round(color.red * 255)), int(round(color.green * 255)), int(round(color.blue * 255))
    )


COL = {
    "DARK_GREY": _hex(NT.DARK_GREY),
    "MED_GREY": _hex(NT.MED_GREY),
    "SOFT_GREY": _hex(NT.SOFT_GREY),
    "ROW_ALT": _hex(NT.ROW_ALT),
    "NOTE_BG": _hex(NT.NOTE_BG),
    "GRID_LINE": _hex(NT.GRID_LINE),
    "INK": _hex(NT.INK),
    "WHITE": "FFFFFF",
}


def rasterize(drawing, out_png, dpi=600):
    """Render a reportlab Drawing to a tight transparent PNG via a one-off PDF."""
    buf = io.BytesIO()
    renderPDF.drawToFile(drawing, buf)
    buf.seek(0)
    doc = pymupdf.open(stream=buf.read(), filetype="pdf")
    page = doc[0]
    zoom = dpi / 72.0
    pix = page.get_pixmap(matrix=pymupdf.Matrix(zoom, zoom), alpha=True)
    pix.save(out_png)
    doc.close()
    return float(drawing.width), float(drawing.height)


# Fixed marks
MARK = {}
w, h = rasterize(NT.motif_dna(42), os.path.join(MARKS_DIR, "motif.png"))
MARK["motif"] = ("motif.png", w, h)
w, h = rasterize(NT._icon_star(11), os.path.join(MARKS_DIR, "star.png"))
MARK["star"] = ("star.png", w, h)
w, h = rasterize(NT._icon_note(11), os.path.join(MARKS_DIR, "note.png"))
MARK["note"] = ("note.png", w, h)


def step_badge_png(n):
    key = f"step{n}"
    if key not in MARK:
        w, h = rasterize(NT._step_badge(n, 16), os.path.join(MARKS_DIR, key + ".png"))
        MARK[key] = (key + ".png", w, h)
    return MARK[key]


def badge_col_width_pt(number, level):
    """Reproduce neet_template.heading()'s badge-column width for `number`."""
    size = {1: 13.5, 2: 11.5, 3: 10.0}[level]
    fs = max(size * 0.46, 6.0)
    pad = fs * 0.42
    text_w = stringWidth(number, NT.FONT_BOLD, fs)
    plate_w = max(size, text_w + 2 * pad)
    BADGE_GUTTER = 5
    return max(1.02 * CM, plate_w + BADGE_GUTTER), size


# ======================================================================================
# PHASE 2 — extract the chapter content VERBATIM by executing the chapter script
# against shims. Semantic blocks are collected into STORY.
# ======================================================================================

class _FakePara:
    def __init__(self, text, style=None, *a, **k):
        self.text = text
        self.style_name = getattr(style, "name", None)


class _FakeImage:
    def __init__(self, path, width=None, height=None, *a, **k):
        self.path = path
        self.width = width
        self.height = height


class _FakeTable:
    def __init__(self, data, colWidths=None, *a, **k):
        self.data = data
        self.colWidths = colWidths
        self.hAlign = None

    def setStyle(self, *a, **k):
        pass


class _FakeTableStyle:
    def __init__(self, *a, **k):
        pass


class _FakeKeep:
    def __init__(self, flowables, *a, **k):
        self.content = flowables


# semantic helper replacements (patched onto the real module before the chapter's
# `from neet_template import ...` binds them)
def _s_title_block(title_text, motif_size=42):
    return [{"k": "title", "text": title_text}]


def _s_heading(number, text, level, has_table=False):
    return {"k": "heading", "number": number, "text": text, "level": level}


def _s_keyterm(text):
    return {"k": "keyterm", "text": text}


def _s_process_flow(steps, cyclic=False):
    return {"k": "flow", "steps": list(steps), "cyclic": cyclic}


def _s_note(text):
    return {"k": "note", "text": text}


def _s_memory(text):
    return {"k": "memory", "text": text}


def _s_data_table(rows, col_widths=None, font_size=9.5):
    return {"k": "table", "rows": rows, "col_widths": col_widths}


def _s_figure(asset_name, caption_text, assets_dir, max_width_cm=15.9):
    return {"k": "figure", "asset": asset_name, "caption": caption_text, "width_cm": max_width_cm}


# patch the real module object
NT.title_block = _s_title_block
NT.heading = _s_heading
NT.keyterm = _s_keyterm
NT.process_flow = _s_process_flow
NT.note = _s_note
NT.memory_aid = _s_memory
NT.data_table = _s_data_table
NT.figure = _s_figure

# install reportlab primitive shims (chapter's local body/b1/b2/figure/figure_row
# build these directly). Real reportlab is no longer needed after Phase 1.
import types  # noqa: E402

_fake_platypus = types.ModuleType("reportlab.platypus")
_fake_platypus.Paragraph = _FakePara
_fake_platypus.Image = _FakeImage
_fake_platypus.Table = _FakeTable
_fake_platypus.TableStyle = _FakeTableStyle
_fake_platypus.KeepTogether = _FakeKeep
sys.modules["reportlab.platypus"] = _fake_platypus

_fake_units = types.ModuleType("reportlab.lib.units")
_fake_units.cm = 1.0  # widths become plain "cm" numbers; only relative sizes matter
sys.modules["reportlab.lib.units"] = _fake_units

# execute the chapter script as a module body (its __main__ guard keeps main() from running)
_g = {"__name__": "ch4_extract", "__file__": CH_SCRIPT}
with open(CH_SCRIPT, "r", encoding="utf-8") as fh:
    code = fh.read()
exec(compile(code, CH_SCRIPT, "exec"), _g)  # noqa: S102

RAW_STORY = _g["story"]


def _normalize(item):
    if isinstance(item, dict):
        return item
    if isinstance(item, _FakePara):
        return {"k": "para", "text": item.text, "style": item.style_name}
    if isinstance(item, _FakeKeep):
        content = item.content
        # figure_row -> KeepTogether(Table([cells]))
        if isinstance(content, _FakeTable):
            row_cells = content.data[0]
            items = []
            for cell in row_cells:
                framed, cap = cell[0], cell[1]
                img = framed.data[0][0]
                asset = os.path.basename(img.path)
                width_cm = img.width if img.width else 6.0
                items.append((asset, cap.text, float(width_cm)))
            return {"k": "figure_row", "items": items}
        # single-figure KeepTogether (not used here, but handle gracefully)
        return {"k": "raw"}
    return {"k": "raw"}


STORY = [_normalize(it) for it in RAW_STORY]

# ======================================================================================
# PHASE 3 — build the .docx
# ======================================================================================
from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402, F401
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.shared import Pt, RGBColor  # noqa: E402

FONT = "Times New Roman"

# style parameters mirrored from neet_template.STYLES
ST = {
    "Title": dict(size=20, bold=True, leading=23, align="center", after=2),
    "H1": dict(size=10.5, bold=True, leading=13, color="WHITE", after=6),
    "H2": dict(size=9.5, bold=True, leading=12, color="WHITE", after=5),
    "H3": dict(size=9.0, bold=True, leading=11.5, color="WHITE", after=4),
    "Body": dict(size=10.8, leading=14.2, after=3),
    "Bullet1": dict(size=10.8, leading=14.2, after=1.5, left=12, hang=8),
    "Bullet2": dict(size=10.5, leading=13.8, after=1.5, left=22, hang=8),
    "Caption": dict(size=9.5, italic=True, leading=12.5, align="center", before=3, after=8),
    "TableHead": dict(size=9.5, bold=True, leading=12, color="WHITE"),
    "TableCell": dict(size=9.5, leading=12),
    "NoteBox": dict(size=10.2, italic=True, leading=13.5),
}
BANNER_BG = {1: "DARK_GREY", 2: "MED_GREY", 3: "SOFT_GREY"}

TOKEN = re.compile(r"(<b>|</b>|<i>|</i>|&bull;|&amp;|&nbsp;|&lt;|&gt;)")


def _set_run(run, size, bold, italic, color_hex):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)


def add_runs(paragraph, text, size, color_hex=None, base_bold=False, base_italic=False):
    bold, ital = base_bold, base_italic
    for part in TOKEN.split(text):
        if part == "<b>":
            bold = True
        elif part == "</b>":
            bold = base_bold
        elif part == "<i>":
            ital = True
        elif part == "</i>":
            ital = base_italic
        elif part == "":
            continue
        else:
            s = (part.replace("&bull;", "\u2022").replace("&amp;", "&")
                 .replace("&nbsp;", "\u00a0").replace("&lt;", "<").replace("&gt;", ">"))
            _set_run(paragraph.add_run(s), size, bold, ital, color_hex)


def _fmt(paragraph, spec):
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(spec["leading"])
    pf.space_before = Pt(spec.get("before", 0))
    pf.space_after = Pt(spec.get("after", 0))
    if spec.get("left"):
        pf.left_indent = Pt(spec["left"])
    if spec.get("hang"):
        pf.first_line_indent = Pt(-spec["hang"])
    a = spec.get("align")
    if a == "center":
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return paragraph


def _shade(el, fill_hex):
    tcpr = el.get_or_add_tcPr() if el.tag.endswith("}tc") else el.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcpr.append(shd)


def _cell_shade(cell, fill_hex):
    _shade(cell._tc, fill_hex)


def _cell_margins(cell, top=0, bottom=0, left=0, right=0):
    tcpr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for name, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right),
                      ("left", left), ("right", right)):
        e = OxmlElement("w:" + name)
        e.set(qn("w:w"), str(int(val * 20)))  # twips
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcpr.append(m)


def _cell_valign(cell, val="center"):
    cell.vertical_alignment = {
        "center": WD_ALIGN_VERTICAL.CENTER, "top": WD_ALIGN_VERTICAL.TOP,
    }[val]


def _set_cell_width(cell, width_pt):
    cell.width = Pt(width_pt)
    tcpr = cell._tc.get_or_add_tcPr()
    tcw = tcpr.find(qn("w:tcW"))
    if tcw is None:
        tcw = OxmlElement("w:tcW")
        tcpr.append(tcw)
    tcw.set(qn("w:w"), str(int(width_pt * 20)))
    tcw.set(qn("w:type"), "dxa")


def _borders(cell, sides, sz_eighths, color_hex):
    tcpr = cell._tc.get_or_add_tcPr()
    tb = tcpr.find(qn("w:tcBorders"))
    if tb is None:
        tb = OxmlElement("w:tcBorders")
        tcpr.append(tb)
    for side in sides:
        e = OxmlElement("w:" + side)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz_eighths))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color_hex)
        tb.append(e)


def _no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + side)
        e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


def _fixed_layout(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    table.autofit = False
    table.allow_autofit = False


def _apply_grid(table, widths_pt):
    """Force the table's column geometry.

    python-docx only writes per-cell w:tcW; Word/LibreOffice under a *fixed*
    layout honour the table-level w:tblGrid + w:tblW instead, and fall back to
    even column distribution when they're missing. Writing them explicitly is
    what makes narrow badge columns and full-width text columns render as laid
    out (fixes badge floating centre-banner and body text shoved to half width).
    Also pins every cell's w:tcW to the same numbers so the two agree.
    """
    twips = [int(round(w * 20)) for w in widths_pt]
    total = sum(twips)
    tbl = table._tbl
    tblPr = tbl.tblPr
    # total table width
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")
    # replace the grid
    old = tbl.find(qn("w:tblGrid"))
    if old is not None:
        tbl.remove(old)
    grid = OxmlElement("w:tblGrid")
    for t in twips:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(t))
        grid.append(gc)
    tblPr.addnext(grid)  # tblGrid must follow tblPr
    # pin each cell
    for row in table.rows:
        for cidx, cell in enumerate(row.cells):
            if cidx >= len(twips):
                continue
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(twips[cidx]))
            tcw.set(qn("w:type"), "dxa")


def _keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def _keep_together(paragraph):
    paragraph.paragraph_format.keep_together = True


def img_display_width_pt(path, max_width_cm):
    from PIL import Image as PILImage
    with PILImage.open(path) as im:
        px_w, _ = im.size
    max_w = min(max_width_cm * CM, FRAME_WIDTH_PT)
    natural_w = px_w / 300.0 * 72.0
    return min(max_w, natural_w)


# ---------- document + page geometry ----------
doc = Document()
sec = doc.sections[0]
sec.page_width = Pt(21.0 * CM)      # A4
sec.page_height = Pt(29.7 * CM)
sec.left_margin = Pt(1.5 * CM)
sec.right_margin = Pt(1.5 * CM)
sec.top_margin = Pt(1.4 * CM)
sec.bottom_margin = Pt(1.4 * CM)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(10.8)
normal._element.get_or_add_rPr().append(OxmlElement("w:rFonts"))
_rf = normal._element.rPr.find(qn("w:rFonts"))
for attr in ("w:ascii", "w:hAnsi", "w:cs"):
    _rf.set(qn(attr), FONT)


def emit_title(block):
    tbl = doc.add_table(rows=1, cols=2)
    _fixed_layout(tbl)
    _no_table_borders(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    motif_col = 42 / 28.35 * CM + 0.1 * CM
    c0, c1 = tbl.rows[0].cells
    _set_cell_width(c0, motif_col)
    _set_cell_width(c1, FRAME_WIDTH_PT - motif_col)
    for c in (c0, c1):
        _cell_margins(c)
        _cell_valign(c, "center")
    name, w, h = MARK["motif"]
    p0 = c0.paragraphs[0]
    p0.add_run().add_picture(os.path.join(MARKS_DIR, name), width=Pt(42))
    p1 = c1.paragraphs[0]
    _fmt(p1, ST["Title"])
    add_runs(p1, block["text"], ST["Title"]["size"], base_bold=True)
    # horizontal rule under the title block
    rule = doc.add_paragraph()
    rp = rule.paragraph_format
    rp.space_before = Pt(4)
    rp.space_after = Pt(8)
    pPr = rule._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "9")            # ~1.1pt
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), COL["DARK_GREY"])
    pbdr.append(bottom)
    pPr.append(pbdr)


def emit_heading(block):
    level = block["level"]
    spec = ST[f"H{level}"]
    badge_w, _size = badge_col_width_pt(block["number"], level)
    tbl = doc.add_table(rows=1, cols=2)
    _fixed_layout(tbl)
    _no_table_borders(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    c0, c1 = tbl.rows[0].cells
    _set_cell_width(c0, badge_w)
    _set_cell_width(c1, FRAME_WIDTH_PT - badge_w)
    _cell_shade(c0, COL["INK"])
    _cell_shade(c1, COL[BANNER_BG[level]])
    _cell_margins(c0, top=2, bottom=3, left=2, right=2)
    _cell_margins(c1, top=2, bottom=3, left=4, right=2)
    _cell_valign(c0, "center")
    _cell_valign(c1, "center")
    p0 = c0.paragraphs[0]
    p0.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p0.paragraph_format.line_spacing = Pt(spec["leading"])
    p0.paragraph_format.space_after = Pt(0)
    _set_run(p0.add_run(block["number"]), spec["size"], True, False, COL["WHITE"])
    p1 = c1.paragraphs[0]
    p1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p1.paragraph_format.line_spacing = Pt(spec["leading"])
    p1.paragraph_format.space_after = Pt(0)
    add_runs(p1, block["text"], spec["size"], color_hex=COL["WHITE"], base_bold=True)
    # bind heading to following text (orphan guard analogue)
    doc.paragraphs[-1] if doc.paragraphs else None


def emit_para(block):
    style_name = block["style"] or "Body"
    spec = ST.get(style_name, ST["Body"])
    p = doc.add_paragraph()
    _fmt(p, spec)
    add_runs(p, block["text"], spec["size"])


def emit_keyterm(block):
    # filled-circle definition marker + body text, matching the 0.5cm marker column
    tbl = doc.add_table(rows=1, cols=2)
    _fixed_layout(tbl)
    _no_table_borders(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    c0, c1 = tbl.rows[0].cells
    _set_cell_width(c0, 0.5 * CM)
    _set_cell_width(c1, FRAME_WIDTH_PT - 0.5 * CM)
    _cell_margins(c0, top=4, right=4)
    _cell_margins(c1, top=0)
    _cell_valign(c0, "top")
    _cell_valign(c1, "top")
    p0 = c0.paragraphs[0]
    p0.paragraph_format.space_after = Pt(0)
    _set_run(p0.add_run("\u25CF"), 7.5, False, False, COL["INK"])
    p1 = c1.paragraphs[0]
    _fmt(p1, ST["Body"])
    add_runs(p1, block["text"], ST["Body"]["size"])


def emit_flow(block):
    steps = block["steps"]
    tbl = doc.add_table(rows=len(steps), cols=2)
    _fixed_layout(tbl)
    _no_table_borders(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, s in enumerate(steps):
        c0, c1 = tbl.rows[i].cells
        _set_cell_width(c0, 0.7 * CM)
        _set_cell_width(c1, FRAME_WIDTH_PT - 0.7 * CM)
        _cell_margins(c0, top=3, bottom=3, right=4)
        _cell_margins(c1, top=3, bottom=3, left=6)
        _cell_valign(c0, "top")
        _cell_valign(c1, "top")
        _borders(c0, ["right"], 6, COL["GRID_LINE"])  # LINEAFTER 0.75pt
        name, w, h = step_badge_png(i + 1)
        p0 = c0.paragraphs[0]
        p0.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_after = Pt(0)
        p0.add_run().add_picture(os.path.join(MARKS_DIR, name), width=Pt(16))
        p1 = c1.paragraphs[0]
        _fmt(p1, ST["Bullet1"])
        p1.paragraph_format.left_indent = Pt(0)
        p1.paragraph_format.first_line_indent = Pt(0)
        add_runs(p1, s, ST["Bullet1"]["size"])


def emit_box(block, kind):
    tbl = doc.add_table(rows=1, cols=1)
    _fixed_layout(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    _set_cell_width(cell, FRAME_WIDTH_PT)
    _cell_shade(cell, COL["NOTE_BG"])
    _cell_margins(cell, top=5, bottom=5, left=6, right=6)
    if kind == "note":
        _borders(cell, ["top", "bottom", "left", "right"], 4, COL["GRID_LINE"])
    else:
        _borders(cell, ["top", "bottom", "left", "right"], 6, COL["GRID_LINE"])
    icon = MARK["note"] if kind == "note" else MARK["star"]
    label = "[NOTE] " if kind == "note" else "[MEMORY AID - not in NCERT] "
    p = cell.paragraphs[0]
    _fmt(p, ST["NoteBox"])
    p.add_run().add_picture(os.path.join(MARKS_DIR, icon[0]), width=Pt(11))
    p.add_run(" ")
    _set_run(p.add_run(label), ST["NoteBox"]["size"], True, True, COL["INK"])
    add_runs(p, block["text"], ST["NoteBox"]["size"], base_italic=True)


def emit_table(block):
    rows = block["rows"]
    col_widths = block["col_widths"]
    ncol = len(rows[0])
    if col_widths:
        tot = sum(col_widths)
        widths = [w / tot * FRAME_WIDTH_PT for w in col_widths]
    else:
        widths = [FRAME_WIDTH_PT / ncol] * ncol
    tbl = doc.add_table(rows=len(rows), cols=ncol)
    _fixed_layout(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for r, row in enumerate(rows):
        is_head = r == 0
        spec = ST["TableHead"] if is_head else ST["TableCell"]
        for cidx, val in enumerate(row):
            cell = tbl.rows[r].cells[cidx]
            _set_cell_width(cell, widths[cidx])
            _borders(cell, ["top", "bottom", "left", "right"], 3, COL["GRID_LINE"])
            _cell_margins(cell, top=3, bottom=3, left=4, right=4)
            _cell_valign(cell, "top")
            if is_head:
                _cell_shade(cell, COL["DARK_GREY"])
            elif r % 2 == 0:
                _cell_shade(cell, COL["ROW_ALT"])
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(spec["leading"])
            p.paragraph_format.space_after = Pt(0)
            add_runs(p, val, spec["size"],
                     color_hex=(COL["WHITE"] if is_head else None),
                     base_bold=is_head)


def _framed_figure_into_cell(cell, asset, width_pt):
    _cell_margins(cell, top=5, bottom=5, left=5, right=5)
    _borders(cell, ["top", "bottom", "left", "right"], 4, COL["GRID_LINE"])
    _cell_valign(cell, "middle" if False else "center")
    p = cell.paragraphs[0]
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(os.path.join(ASSETS, asset), width=Pt(width_pt))


def emit_figure(block):
    asset = block["asset"]
    width_pt = img_display_width_pt(os.path.join(ASSETS, asset), block["width_cm"])
    outer = width_pt + 10
    tbl = doc.add_table(rows=1, cols=1)
    _fixed_layout(tbl)
    _no_table_borders(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    _set_cell_width(cell, outer)
    _framed_figure_into_cell(cell, asset, width_pt)
    cap = doc.add_paragraph()
    _fmt(cap, ST["Caption"])
    _keep_together(cap)
    add_runs(cap, block["caption"], ST["Caption"]["size"], base_italic=True)


def emit_figure_row(block):
    items = block["items"]
    widths_pt = [img_display_width_pt(os.path.join(ASSETS, a), wc) for a, _c, wc in items]
    pad = 6
    tbl = doc.add_table(rows=1, cols=len(items))
    _fixed_layout(tbl)
    _no_table_borders(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (asset, caption, _wc) in enumerate(items):
        wpt = widths_pt[idx]
        cell = tbl.rows[0].cells[idx]
        _set_cell_width(cell, wpt + 10 + 2 * pad)
        _cell_valign(cell, "top")
        _cell_margins(cell, left=pad, right=pad)
        # inner framed image
        inner = cell.add_table(rows=1, cols=1)
        _fixed_layout(inner)
        _no_table_borders(inner)
        inner.alignment = WD_TABLE_ALIGNMENT.CENTER
        icell = inner.rows[0].cells[0]
        _set_cell_width(icell, wpt + 10)
        _framed_figure_into_cell(icell, asset, wpt)
        capp = cell.add_paragraph()
        _fmt(capp, ST["Caption"])
        add_runs(capp, caption, ST["Caption"]["size"], base_italic=True)
    # clear the default empty first paragraph inside each cell is left as spacing


EMIT = {
    "title": emit_title,
    "heading": emit_heading,
    "para": emit_para,
    "keyterm": emit_keyterm,
    "flow": emit_flow,
    "note": lambda b: emit_box(b, "note"),
    "memory": lambda b: emit_box(b, "memory"),
    "table": emit_table,
    "figure": emit_figure,
    "figure_row": emit_figure_row,
}


def main():
    for block in STORY:
        fn = EMIT.get(block["k"])
        if fn:
            fn(block)
    doc.save(OUT_DOCX)
    size_kb = os.path.getsize(OUT_DOCX) / 1024
    print(f"Built {OUT_DOCX} ({size_kb:.0f} KB) from {len(STORY)} blocks")
    return 0


if __name__ == "__main__":
    sys.exit(main())
