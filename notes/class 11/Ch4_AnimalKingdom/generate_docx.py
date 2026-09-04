import sys
import os
import re
from xml.sax.saxutils import unescape

import docx
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

sys.path.append('notes/class 11/Ch4_AnimalKingdom')
import Ch4_AnimalKingdom
from reportlab.platypus import KeepTogether, Table, Paragraph, Image, HRFlowable
from reportlab.graphics.shapes import Drawing, String, Group

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color='E5E7EB', sz='4', val='single'):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:left w:val="none"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:right w:val="none"/>
            <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def set_box_borders(table, color='1F2937', sz='12', val='single', left_only=False):
    tblPr = table._tbl.tblPr
    if left_only:
        borders = parse_xml(f'''
            <w:tblBorders {nsdecls("w")}>
                <w:top w:val="none"/>
                <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:bottom w:val="none"/>
                <w:right w:val="none"/>
                <w:insideH w:val="none"/>
                <w:insideV w:val="none"/>
            </w:tblBorders>
        ''')
    else:
        borders = parse_xml(f'''
            <w:tblBorders {nsdecls("w")}>
                <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:insideH w:val="none"/>
                <w:insideV w:val="none"/>
            </w:tblBorders>
        ''')
    tblPr.append(borders)

def add_formatted_runs(paragraph, text, base_font_size=10.5, base_color='111827', default_bold=False, default_italic=False):
    text = text.replace('&bull;', '• ').replace('&nbsp;', ' ')
    tokens = re.split(r'(</?[bi]>|<font[^>]*>|</font>)', text)
    is_bold = default_bold
    is_italic = default_italic
    
    for tok in tokens:
        if not tok:
            continue
        tok_lower = tok.lower()
        if tok_lower == '<b>':
            is_bold = True
        elif tok_lower == '</b>':
            is_bold = False
        elif tok_lower == '<i>':
            is_italic = True
        elif tok_lower == '</i>':
            is_italic = False
        elif tok_lower.startswith('<font') or tok_lower == '</font>':
            continue
        else:
            clean_str = unescape(tok, {'&amp;': '&', '&lt;': '<', '&gt;': '>', '&quot;': '"', '&apos;': "'"})
            run = paragraph.add_run(clean_str)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(base_font_size)
            run.bold = is_bold
            run.italic = is_italic
            if base_color:
                r = int(base_color[0:2], 16)
                g = int(base_color[2:4], 16)
                b = int(base_color[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)

def extract_heading_info(tbl):
    cells = tbl._cellvalues
    badge_text = ''
    heading_text = ''
    c0 = cells[0][0]
    if isinstance(c0, Drawing):
        for shape in c0.contents:
            if isinstance(shape, String):
                badge_text = shape.text
            elif isinstance(shape, Group):
                for sub in shape.contents:
                    if isinstance(sub, String):
                        badge_text = sub.text
    c1 = cells[0][1]
    if isinstance(c1, Paragraph):
        heading_text = c1.text
    return badge_text, heading_text

def build_docx(out_path):
    doc = docx.Document()
    
    # Page setup matching neet_template.py: A4, 1.5cm left/right, 1.4cm top/bottom
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    
    # Iterate over Ch4_AnimalKingdom.story
    story = Ch4_AnimalKingdom.story
    
    for i, el in enumerate(story):
        # 1. Title block (Item 0)
        if i == 0:
            p_chap = doc.add_paragraph()
            p_chap.paragraph_format.space_before = Pt(0)
            p_chap.paragraph_format.space_after = Pt(2)
            r_chap = p_chap.add_run("CHAPTER 4")
            r_chap.font.name = "Times New Roman"
            r_chap.font.size = Pt(11)
            r_chap.font.bold = True
            r_chap.font.color.rgb = RGBColor(107, 114, 128) # SOFT_GREY
            
            p_title = doc.add_paragraph()
            p_title.paragraph_format.space_before = Pt(0)
            p_title.paragraph_format.space_after = Pt(6)
            r_title = p_title.add_run("Animal Kingdom")
            r_title.font.name = "Times New Roman"
            r_title.font.size = Pt(24)
            r_title.font.bold = True
            r_title.font.color.rgb = RGBColor(17, 24, 39) # INK
            continue
            
        # 2. HRFlowable (Item 1)
        if isinstance(el, HRFlowable):
            p_hr = doc.add_paragraph()
            p_hr.paragraph_format.space_before = Pt(0)
            p_hr.paragraph_format.space_after = Pt(10)
            p_hr_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="8" w:space="1" w:color="1F2937"/></w:pBdr>')
            p_hr._p.get_or_add_pPr().append(p_hr_border)
            continue
            
        # 3. Headings (wrapped in KeepTogether with Table)
        if isinstance(el, KeepTogether):
            content = el._content
            
            # Check if this is a heading banner
            is_heading = any(isinstance(sub, Table) and len(sub._cellvalues) == 1 and len(sub._cellvalues[0]) == 2 and isinstance(sub._cellvalues[0][0], Drawing) for sub in content)
            if is_heading:
                tbl = [sub for sub in content if isinstance(sub, Table)][0]
                badge_text, heading_text = extract_heading_info(tbl)
                
                # Determine level
                if badge_text in ['4', 'Summary', 'Exercises']:
                    level = 1
                elif '.' in badge_text:
                    dots = badge_text.count('.')
                    if dots == 1:
                        level = 1
                    elif dots == 2:
                        level = 2
                    else:
                        level = 3
                else:
                    level = 2
                    
                # Create banner table
                banner = doc.add_table(rows=1, cols=2)
                banner.alignment = WD_TABLE_ALIGNMENT.LEFT
                banner.autofit = False
                
                badge_w = 1.3 if len(badge_text) <= 5 else 2.4
                banner.columns[0].width = Cm(badge_w)
                banner.columns[1].width = Cm(18.0 - badge_w)
                
                # Badge cell
                c0 = banner.rows[0].cells[0]
                c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_background(c0, '1F2937') # DARK_GREY
                set_cell_margins(c0, top=60, bottom=60, left=80, right=80)
                p0 = c0.paragraphs[0]
                p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p0.paragraph_format.space_before = Pt(0)
                p0.paragraph_format.space_after = Pt(0)
                r0 = p0.add_run(badge_text)
                r0.font.name = 'Times New Roman'
                r0.font.size = Pt(8.5 if len(badge_text) > 4 else 9.5)
                r0.font.bold = True
                r0.font.color.rgb = RGBColor(255, 255, 255)
                
                # Text cell
                c1 = banner.rows[0].cells[1]
                c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_margins(c1, top=60, bottom=60, left=120, right=80)
                p1 = c1.paragraphs[0]
                p1.paragraph_format.space_before = Pt(0)
                p1.paragraph_format.space_after = Pt(0)
                font_sz = {1: 13.5, 2: 11.5, 3: 10.5}.get(level, 11.0)
                add_formatted_runs(p1, heading_text, base_font_size=font_sz, base_color='111827', default_bold=True)
                
                # Border
                tblPr = banner._tbl.tblPr
                if level == 1:
                    borders = parse_xml(f'''
                        <w:tblBorders {nsdecls("w")}>
                            <w:top w:val="none"/><w:left w:val="none"/>
                            <w:bottom w:val="single" w:sz="8" w:space="0" w:color="1F2937"/>
                            <w:right w:val="none"/><w:insideH w:val="none"/><w:insideV w:val="none"/>
                        </w:tblBorders>
                    ''')
                else:
                    borders = parse_xml(f'''
                        <w:tblBorders {nsdecls("w")}>
                            <w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/>
                            <w:right w:val="none"/><w:insideH w:val="none"/><w:insideV w:val="none"/>
                        </w:tblBorders>
                    ''')
                tblPr.append(borders)
                
                # Small spacer after heading banner
                p_spacer = doc.add_paragraph()
                p_spacer.paragraph_format.space_before = Pt(0)
                p_spacer.paragraph_format.space_after = Pt(4)
                continue
                
            # Check if single figure (KeepTogether([Table(1x1 Image), Paragraph(Caption)]))
            if len(content) == 2 and isinstance(content[0], Table) and isinstance(content[1], Paragraph):
                tbl_img = content[0]
                cap_para = content[1]
                img = tbl_img._cellvalues[0][0]
                img_path = getattr(img, 'filename', getattr(img, '_file', ''))
                img_w_cm = img._width / 28.3465
                
                # Add image centered
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(6)
                p_img.paragraph_format.space_after = Pt(2)
                p_img.paragraph_format.keep_with_next = True
                
                if os.path.exists(img_path):
                    run_img = p_img.add_run()
                    run_img.add_picture(img_path, width=Cm(min(img_w_cm, 16.0)))
                
                # Add caption centered
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.space_before = Pt(2)
                p_cap.paragraph_format.space_after = Pt(8)
                add_formatted_runs(p_cap, cap_para.text, base_font_size=9.0, base_color='4B5563', default_italic=True)
                continue
                
            # Check if side-by-side figures (KeepTogether([Table 1x2 with [Table, Paragraph]]))
            if len(content) == 1 and isinstance(content[0], Table):
                tbl_row = content[0]
                if len(tbl_row._cellvalues) == 1 and len(tbl_row._cellvalues[0]) == 2:
                    side_tbl = doc.add_table(rows=1, cols=2)
                    side_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    side_tbl.autofit = False
                    side_tbl.columns[0].width = Cm(9.0)
                    side_tbl.columns[1].width = Cm(9.0)
                    
                    for ci, cell_items in enumerate(tbl_row._cellvalues[0]):
                        c = side_tbl.rows[0].cells[ci]
                        c.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                        set_cell_margins(c, top=60, bottom=60, left=60, right=60)
                        
                        img_sub_tbl = cell_items[0]
                        cap_sub_p = cell_items[1]
                        
                        img_obj = img_sub_tbl._cellvalues[0][0]
                        sub_img_path = getattr(img_obj, 'filename', getattr(img_obj, '_file', ''))
                        sub_w_cm = img_obj._width / 28.3465
                        
                        p_sub_img = c.paragraphs[0]
                        p_sub_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_sub_img.paragraph_format.space_before = Pt(0)
                        p_sub_img.paragraph_format.space_after = Pt(2)
                        p_sub_img.paragraph_format.keep_with_next = True
                        if os.path.exists(sub_img_path):
                            r_s = p_sub_img.add_run()
                            r_s.add_picture(sub_img_path, width=Cm(min(sub_w_cm, 8.5)))
                            
                        p_sub_cap = c.add_paragraph()
                        p_sub_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_sub_cap.paragraph_format.space_before = Pt(2)
                        p_sub_cap.paragraph_format.space_after = Pt(4)
                        add_formatted_runs(p_sub_cap, cap_sub_p.text, base_font_size=8.5, base_color='4B5563', default_italic=True)
                        
                    # Remove borders from side-by-side table
                    tblPr = side_tbl._tbl.tblPr
                    borders = parse_xml(f'''
                        <w:tblBorders {nsdecls("w")}>
                            <w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/>
                            <w:right w:val="none"/><w:insideH w:val="none"/><w:insideV w:val="none"/>
                        </w:tblBorders>
                    ''')
                    tblPr.append(borders)
                    
                    p_sp = doc.add_paragraph()
                    p_sp.paragraph_format.space_before = Pt(0)
                    p_sp.paragraph_format.space_after = Pt(6)
                    continue

        # 4. Paragraphs
        if isinstance(el, Paragraph):
            style_name = el.style.name
            p = doc.add_paragraph()
            
            if 'Bullet' in style_name:
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.15
                add_formatted_runs(p, el.text, base_font_size=10.5, base_color='111827')
            elif style_name == 'Body':
                p.paragraph_format.left_indent = Cm(0)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.15
                add_formatted_runs(p, el.text, base_font_size=10.5, base_color='111827')
            else:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(3)
                add_formatted_runs(p, el.text, base_font_size=10.5, base_color='111827')
            continue
            
        # 5. Tables
        if isinstance(el, Table):
            cells = el._cellvalues
            r_count = len(cells)
            c_count = len(cells[0]) if r_count > 0 else 0
            
            # 5a. Keyterm (1x2 table with Drawing in 0,0)
            if r_count == 1 and c_count == 2 and isinstance(cells[0][0], Drawing):
                p_kt = doc.add_paragraph()
                p_kt.paragraph_format.left_indent = Cm(0.5)
                p_kt.paragraph_format.space_before = Pt(1)
                p_kt.paragraph_format.space_after = Pt(3)
                r_bullet = p_kt.add_run("•  ")
                r_bullet.font.name = "Times New Roman"
                r_bullet.font.size = Pt(10.5)
                r_bullet.font.bold = True
                r_bullet.font.color.rgb = RGBColor(31, 41, 55)
                add_formatted_runs(p_kt, cells[0][1].text, base_font_size=10.5, base_color='111827')
                continue
                
            # 5b. Note / Memory Aid (1x1 table with nested Table)
            if r_count == 1 and c_count == 1 and isinstance(cells[0][0], Table):
                nested = cells[0][0]
                box_tbl = doc.add_table(rows=1, cols=1)
                box_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                box_tbl.autofit = False
                box_tbl.columns[0].width = Cm(18.0)
                
                box_cell = box_tbl.rows[0].cells[0]
                set_cell_background(box_cell, 'F9FAFB') # NOTE_BG
                set_cell_margins(box_cell, top=100, bottom=100, left=150, right=150)
                
                # Check text to see if NOTE or MEMORY AID
                para_obj = None
                for nr in nested._cellvalues:
                    for nc in nr:
                        if isinstance(nc, Paragraph):
                            para_obj = nc
                            break
                            
                p_box = box_cell.paragraphs[0]
                p_box.paragraph_format.space_before = Pt(0)
                p_box.paragraph_format.space_after = Pt(0)
                p_box.paragraph_format.line_spacing = 1.15
                
                is_mem = 'MEMORY AID' in (para_obj.text if para_obj else '')
                border_color = '4B5563' if is_mem else '1F2937'
                border_val = 'dashed' if is_mem else 'single'
                
                set_box_borders(box_tbl, color=border_color, sz='16', val=border_val, left_only=True)
                
                if para_obj:
                    add_formatted_runs(p_box, para_obj.text, base_font_size=10.0, base_color='1F2937')
                    
                p_sp = doc.add_paragraph()
                p_sp.paragraph_format.space_before = Pt(0)
                p_sp.paragraph_format.space_after = Pt(6)
                continue
                
            # 5c. Process flow (3x2 table with Drawings in col 0)
            if r_count == 3 and c_count == 2 and isinstance(cells[0][0], Drawing):
                flow_tbl = doc.add_table(rows=3, cols=2)
                flow_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                flow_tbl.autofit = False
                flow_tbl.columns[0].width = Cm(1.2)
                flow_tbl.columns[1].width = Cm(16.8)
                
                for s_idx in range(3):
                    # Step badge cell
                    step_c0 = flow_tbl.rows[s_idx].cells[0]
                    step_c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    set_cell_background(step_c0, '1F2937')
                    set_cell_margins(step_c0, top=60, bottom=60, left=60, right=60)
                    sp0 = step_c0.paragraphs[0]
                    sp0.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    sr0 = sp0.add_run(f"{s_idx + 1}")
                    sr0.font.name = "Times New Roman"
                    sr0.font.size = Pt(10)
                    sr0.font.bold = True
                    sr0.font.color.rgb = RGBColor(255, 255, 255)
                    
                    # Step text cell
                    step_c1 = flow_tbl.rows[s_idx].cells[1]
                    step_c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    set_cell_background(step_c1, 'F3F4F6' if s_idx % 2 == 1 else 'FFFFFF')
                    set_cell_margins(step_c1, top=60, bottom=60, left=100, right=80)
                    sp1 = step_c1.paragraphs[0]
                    sp1.paragraph_format.space_before = Pt(0)
                    sp1.paragraph_format.space_after = Pt(0)
                    add_formatted_runs(sp1, cells[s_idx][1].text, base_font_size=10.0, base_color='111827')
                    
                set_table_borders(flow_tbl, color='E5E7EB', sz='4', val='single')
                
                p_sp = doc.add_paragraph()
                p_sp.paragraph_format.space_before = Pt(0)
                p_sp.paragraph_format.space_after = Pt(6)
                continue
                
            # 5d. Real Data Tables (Table 4.1, Table 4.2)
            word_tbl = doc.add_table(rows=r_count, cols=c_count)
            word_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            word_tbl.autofit = False
            
            # Compute total width available: 18.0 cm
            total_w = 18.0
            col_w = total_w / c_count
            for col in word_tbl.columns:
                col.width = Cm(col_w)
                
            for ri, row in enumerate(cells):
                w_row = word_tbl.rows[ri]
                # Repeat header row
                if ri == 0:
                    trPr = w_row._tr.get_or_add_trPr()
                    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
                # Prevent row split
                trPr = w_row._tr.get_or_add_trPr()
                trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
                
                for ci, cell_val in enumerate(row):
                    w_cell = w_row.cells[ci]
                    w_cell.width = Cm(col_w)
                    w_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    # Background shading: row 0 dark grey, odd rows row_alt, even white
                    if ri == 0:
                        set_cell_background(w_cell, '1F2937') # DARK_GREY
                        txt_color = 'FFFFFF'
                        is_b = True
                    else:
                        bg_color = 'F3F4F6' if ri % 2 == 1 else 'FFFFFF'
                        set_cell_background(w_cell, bg_color)
                        txt_color = '111827'
                        is_b = False
                        
                    set_cell_margins(w_cell, top=60, bottom=60, left=60, right=60)
                    
                    p_cell = w_cell.paragraphs[0]
                    p_cell.paragraph_format.space_before = Pt(0)
                    p_cell.paragraph_format.space_after = Pt(0)
                    p_cell.paragraph_format.line_spacing = 1.05
                    
                    if ri == 0:
                        p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                    cell_text = cell_val.text if isinstance(cell_val, Paragraph) else str(cell_val)
                    font_sz = 8.5 if c_count > 4 else 9.5
                    add_formatted_runs(p_cell, cell_text, base_font_size=font_sz, base_color=txt_color, default_bold=is_b)
                    
            set_table_borders(word_tbl, color='D1D5DB', sz='4', val='single')
            
            p_sp = doc.add_paragraph()
            p_sp.paragraph_format.space_before = Pt(0)
            p_sp.paragraph_format.space_after = Pt(6)
            continue

    doc.save(out_path)
    print(f"Successfully generated clean native DOCX at: {out_path} ({os.path.getsize(out_path)} bytes)")

if __name__ == '__main__':
    out_file = 'notes/class 11/Ch4_AnimalKingdom/Ch4_AnimalKingdom.docx'
    if len(sys.argv) > 1:
        out_file = sys.argv[1]
    build_docx(out_file)
